import docx
import json
import os
import re

# Folder with the docs
folder = r"d:\smira progickt\New folder"

# Months mapping
months_files = {
    "1": {"type1": "احصائية_تفصيلية_لكل_طبيب_يناير_2026_.docx", "type2": "شهر 1 2026.docx"},
    "2": {"type1": "احصائية تفصيلية لكل طبيب شباط2026.docx", "type2": "شهر 2 2026.docx"},
    "3": {"type1": "احصائية_تفصيلية_لكل_طبيب_اذار2026_Copy.docx", "type2": "شهر 3  عتبة2026.docx"},
    "4": {"type1": "احصائية_تفصيلية_لكل_طبيب_نيسان2026_Copy_Copy.docx", "type2": "شهر4عتبة2026.docx"},
    "5": {"type1": "احصائية_تفصيلية_لكل_طبيب_ايار_2026_.docx", "type2": "شهر5عتبة2026.docx"},
    "6": {"type1": "احصائية_تفصيلية_لكل_طبيب_حزيران2026_Copy_Copy.docx", "type2": "شهر6عتبة2026.docx"}
}

def clean_text(text):
    return text.strip().replace('\n', ' ')

def parse_type2(file_path):
    print(f"Parsing {file_path}")
    doc = docx.Document(file_path)
    data = {
        "doctors_visits": [],
        "governorates_visits": [],
        "countries_visits": [],
        "eye_tests": [],
        "lab_tests": []
    }
    
    for idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2: continue
        
        headers = [clean_text(c.text) for c in rows[0].cells]
        
        # 1. Doctors Visits
        if "اسم الطبيب" in headers and "عدد المراجعين" in headers:
            for r in rows[1:]:
                cells = [clean_text(c.text) for c in r.cells]
                if len(cells) >= 3 and cells[1] and cells[2].isdigit():
                    data["doctors_visits"].append({"doctor": cells[1], "count": int(cells[2])})
                    
        # 2. Governorates Visits
        if "المحافظة" in headers and len(rows) >= 2:
            govs = [clean_text(c.text) for c in rows[0].cells]
            counts = [clean_text(c.text) for c in rows[1].cells]
            if "المحافظة" in govs[0]:
                for i in range(1, len(govs)):
                    if govs[i] != "المجموع" and counts[i].isdigit():
                        data["governorates_visits"].append({"governorate": govs[i], "count": int(counts[i])})
                        
        # 3. Countries Visits
        if "البلد" in headers and len(rows) >= 2:
            countries = [clean_text(c.text) for c in rows[0].cells]
            counts = [clean_text(c.text) for c in rows[1].cells]
            if "البلد" in countries[0]:
                for i in range(1, len(countries)):
                    if countries[i] != "المجموع" and counts[i].isdigit():
                        data["countries_visits"].append({"country": countries[i], "count": int(counts[i])})
                        
        # 4. Eye Tests
        if len(headers) == 2 and ("فحص" in headers[0] or "فحص" in rows[1].cells[0].text):
            for r in rows:
                cells = [clean_text(c.text) for c in r.cells]
                if len(cells) == 2 and cells[1].isdigit():
                    data["eye_tests"].append({"test": cells[0], "count": int(cells[1])})
                    
        # 5. Lab Tests (Usually in a horizontal format with headers)
        if any("فحص" in h for h in headers) and len(rows) >= 2 and len(headers) > 2:
            counts = [clean_text(c.text) for c in rows[1].cells]
            for i in range(len(headers)):
                if counts[i].isdigit():
                    data["lab_tests"].append({"test": headers[i], "count": int(counts[i])})
    
    return data

def parse_type1(file_path):
    # For now just list tables for surgeries
    print(f"Parsing {file_path}")
    doc = docx.Document(file_path)
    surgeries = []
    
    for idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2: continue
        headers = [clean_text(c.text) for c in rows[0].cells]
        if "اسم العملية" in headers and "تصنيف العملية" in headers and "عددها" in headers:
            for r in rows[1:]:
                cells = [clean_text(c.text) for c in r.cells]
                if len(cells) >= 3 and cells[0].isdigit():
                    surgeries.append({
                        "operation": cells[2],
                        "classification": cells[1],
                        "count": int(cells[0])
                    })
    return surgeries

results = {}
for month, files in months_files.items():
    print(f"\n--- Processing Month {month} ---")
    type2_data = parse_type2(os.path.join(folder, files["type2"]))
    type1_data = parse_type1(os.path.join(folder, files["type1"]))
    
    results[month] = {
        "visits": type2_data,
        "surgeries": type1_data
    }

with open(r"d:\smira progickt\statistics\database\data\extracted_statistics.json", "w", encoding="utf-8") as f:
    json.dump(results, f, ensure_ascii=False, indent=4)

print("\nData successfully extracted to database/data/extracted_statistics.json")
